import re
import asyncio
import base64
import random
import time
from typing import Any, Optional, Union

import psutil
from dotenv import load_dotenv
from fastapi import (
    BackgroundTasks,
    Depends,
    FastAPI,
    File,
    Request,
    UploadFile,
    WebSocket,
    WebSocketDisconnect,
)
from fastapi.exceptions import RequestValidationError
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response

# run_in_threadpool 已移除：handwrite() 返回惰性生成器（map 对象），
# 真正的 CPU 密集渲染在后续 for 循环消费生成器时才发生，
# 而 generate_handwriting_impl 整体运行在 BackgroundTask 里，HTTP 请求已秒回，
# 所以 run_in_threadpool 在此场景下无法真正释放事件循环。
from handright import Template, handwrite

# from threading import Thread
from PIL import Image, ImageDraw, ImageFont, ImageEnhance, ImageFilter, ImageOps

load_dotenv()
import gc
import io
import logging
import os
import shutil
import tempfile
import json
import urllib.request
import urllib.error
from uuid import uuid4

import PyPDF2

# 文件模块
from docx import Document

# 图片处理模块
from identify import identify_distance
from pdf import generate_pdf
from werkzeug.utils import secure_filename


# 安全文件删除函数
def safe_remove_directory(directory_path, max_retries=5):
    """安全删除目录，带重试机制和更强的文件处理"""
    if not os.path.exists(directory_path):
        return True

    for attempt in range(max_retries):
        try:
            # 强制垃圾回收，释放可能的文件句柄
            gc.collect()
            # 等待更长时间让系统释放文件句柄
            time.sleep(0.2 * (attempt + 1))  # 递增等待时间

            # 递归删除所有文件和子目录
            deleted_files = []
            failed_files = []

            for root, dirs, files in os.walk(directory_path, topdown=False):
                # 删除文件
                for file in files:
                    file_path = os.path.join(root, file)
                    if safe_remove_single_file(file_path, max_retries=2):
                        deleted_files.append(file_path)
                    else:
                        failed_files.append(file_path)

                # 删除空目录
                for dir in dirs:
                    dir_path = os.path.join(root, dir)
                    try:
                        if os.path.exists(dir_path) and not os.listdir(dir_path):
                            os.rmdir(dir_path)
                    except Exception as dir_e:
                        logger.warning(
                            f"Failed to delete subdirectory {dir_path}: {dir_e}"
                        )

            # 如果有文件删除失败，记录但继续尝试删除目录
            if failed_files:
                logger.warning(
                    f"Failed to delete {len(failed_files)} files: {failed_files[:3]}..."
                )

            # 最后尝试删除根目录
            if os.path.exists(directory_path):
                try:
                    os.rmdir(directory_path)
                    logger.info(
                        f"Successfully deleted temp directory: {directory_path}"
                    )
                    return True
                except OSError as e:
                    if e.errno == 145:  # 目录不为空
                        # 列出剩余文件
                        remaining_files = []
                        try:
                            for root, dirs, files in os.walk(directory_path):
                                remaining_files.extend(
                                    [os.path.join(root, f) for f in files]
                                )
                        except:
                            pass
                        logger.warning(
                            f"Directory not empty, remaining files: {remaining_files[:5]}"
                        )
                    raise

        except Exception as e:
            logger.warning(
                f"Attempt {attempt + 1} to delete {directory_path} failed: {e}"
            )
            if attempt < max_retries - 1:
                time.sleep(1.0 * (attempt + 1))  # 递增等待时间
            else:
                logger.error(
                    f"Failed to delete temp directory after {max_retries} attempts: {directory_path}"
                )
                # 最后尝试：标记目录为稍后清理
                try:
                    cleanup_marker = os.path.join(directory_path, ".cleanup_later")
                    with open(cleanup_marker, "w") as f:
                        f.write(f"Failed to delete at {time.time()}")
                    logger.info(f"Marked directory for later cleanup: {directory_path}")
                except:
                    pass
    return False


def safe_remove_single_file(file_path, max_retries=3):
    """安全删除单个文件"""
    if not os.path.exists(file_path):
        return True

    for attempt in range(max_retries):
        try:
            # 确保文件不是只读
            os.chmod(file_path, 0o777)

            # 强制垃圾回收
            gc.collect()
            time.sleep(0.1)

            # 尝试删除文件
            os.remove(file_path)
            return True

        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(0.3 * (attempt + 1))
            else:
                logger.warning(f"Failed to delete file {file_path}: {e}")
    return False


def safe_remove_file(file_path, max_retries=3):
    """安全删除文件，带重试机制"""
    result = safe_remove_single_file(file_path, max_retries)
    if result:
        logger.info(f"Successfully deleted file: {file_path}")
    else:
        logger.error(f"Failed to delete file after {max_retries} attempts: {file_path}")
    return result


def safe_save_and_close_image(image, image_path):
    """安全保存并关闭图片，确保文件句柄被释放"""
    try:
        # 保存图片
        image.save(image_path)

        # 如果图片对象有 close 方法，调用它
        # if hasattr(image, "close"):
        #     image.close()

        # # 强制垃圾回收
        # gc.collect()

        # # 等待一小段时间确保文件句柄被释放
        # time.sleep(0.1)

        return True
    except Exception as e:
        logger.error(f"Failed to save image {image_path}: {e}")
        return False


def cleanup_marked_directories():
    """清理项目内标记为稍后清理的目录"""
    project_temp_base = "./temp"

    # 确保项目临时目录存在
    os.makedirs(project_temp_base, exist_ok=True)

    try:
        for item in os.listdir(project_temp_base):
            item_path = os.path.join(project_temp_base, item)
            if os.path.isdir(item_path) and item.startswith("tmp"):
                cleanup_marker = os.path.join(item_path, ".cleanup_later")
                if os.path.exists(cleanup_marker):
                    try:
                        # 检查标记时间，如果超过1小时则尝试清理
                        with open(cleanup_marker, "r") as f:
                            content = f.read()
                            if "Failed to delete at" in content:
                                timestamp = float(
                                    content.split("Failed to delete at ")[1]
                                )
                                if time.time() - timestamp > 3600:  # 1小时后
                                    logger.info(
                                        f"Attempting to cleanup marked directory: {item_path}"
                                    )
                                    if safe_remove_directory(item_path, max_retries=2):
                                        logger.info(
                                            f"Successfully cleaned up marked directory: {item_path}"
                                        )
                    except Exception as e:
                        logger.warning(
                            f"Failed to cleanup marked directory {item_path}: {e}"
                        )
    except Exception as e:
        logger.warning(f"Error during cleanup of marked directories: {e}")


# 装饰器 7.15
from functools import wraps

# 定时清理文件 10.28
import schedule_clean

# sentry 错误报告7.7
import sentry_sdk
from sentry_sdk.integrations.fastapi import FastApiIntegration
from sentry_sdk.integrations.starlette import StarletteIntegration

# 限制请求速率 7.9
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address
from task_store import cleanup_expired as cleanup_expired_generation_tasks
from task_store import get_active_task_count as get_generation_active_task_count
from task_store import get_queue_metrics as get_generation_queue_metrics
from task_store import get_task as get_generation_task
from task_store import pop_task as pop_generation_task
from task_store import read_result_file
from task_store import set_task as set_generation_task
from task_types import (
    GenerateHandwritingParams,
    GenerationTask,
    form_dependency_from_model,
)

# 获取环境变量
mysql_host = os.getenv("MYSQL_HOST", "db")
enable_user_auth = os.getenv("ENABLE_USER_AUTH", "false")
# 获取当前路径
current_path = os.getcwd()
# 创建一个子文件夹用于存储输出的图片
output_path = os.path.join(current_path, "output")
# 如果子文件夹不存在，就创建它
if not os.path.exists(output_path):
    os.makedirs(output_path)
directory = ["./textfileprocess", "imagefileprocess"]
for directory in directory:
    if not os.path.exists(directory):
        os.makedirs(directory)

font_assets_dir = os.getenv("FONT_ASSETS_DIR", "./font_assets")
font_assets_bundled_dir = os.getenv("FONT_ASSETS_BUNDLED_DIR", "./font_assets")

def sync_font_assets(source_dir, target_dir):
    if os.path.abspath(source_dir) == os.path.abspath(target_dir):
        return
    if not os.path.isdir(source_dir) or not os.path.isdir(target_dir):
        return
    for filename in os.listdir(source_dir):
        if not filename.lower().endswith((".ttf", ".otf")):
            continue
        source_path = os.path.join(source_dir, filename)
        target_path = os.path.join(target_dir, filename)
        if os.path.isfile(source_path) and not os.path.exists(target_path):
            shutil.copy2(source_path, target_path)

os.makedirs(font_assets_dir, exist_ok=True)
sync_font_assets(font_assets_bundled_dir, font_assets_dir)

font_file_names = [
    f
    for f in os.listdir(font_assets_dir)
    if os.path.isfile(os.path.join(font_assets_dir, f)) and f.lower().endswith((".ttf", ".otf"))
]


def refresh_font_file_names():
    """Refresh and repair font list at runtime."""
    global font_file_names
    try:
        os.makedirs(font_assets_dir, exist_ok=True)
        sync_font_assets(font_assets_bundled_dir, font_assets_dir)
        # Extra fallback: copy from backend/font_assets when env path is bad or empty.
        backend_bundle = os.path.join(os.path.dirname(os.path.abspath(__file__)), "font_assets")
        sync_font_assets(backend_bundle, font_assets_dir)
        font_file_names = [
            f for f in os.listdir(font_assets_dir)
            if os.path.isfile(os.path.join(font_assets_dir, f)) and f.lower().endswith((".ttf", ".otf"))
        ]
        font_file_names = sorted(font_file_names, key=lambda x: (0 if x.lower().startswith("sukoya") else 1, x.lower()))
    except Exception as exc:
        logger.warning(f"refresh_font_file_names failed: {exc}")
        font_file_names = []
    return font_file_names

refresh_font_file_names()

SYSTEM_FONT_AUTO = "__system__:auto"
PROFILE_SEUOMI_VIRTUAL = "__profile__:seuomi"

def get_system_font_candidates():
    """Local fallback only. No bundled third-party fonts."""
    candidates = [
        # Windows
        r"C:\Windows\Fonts\simkai.ttf",
        r"C:\Windows\Fonts\kaiu.ttf",
        r"C:\Windows\Fonts\STKAITI.TTF",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyh.ttf",
        # macOS
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Supplemental/Kaiti.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        # Linux common CJK
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    ]
    return candidates

def resolve_system_font_path():
    for path in get_system_font_candidates():
        if os.path.exists(path):
            return path
    return None


def normalize_handwriting_profile(profile):
    value = str(profile or "sukoya").strip().lower()
    return "seuomi" if value == "seuomi" else "sukoya"

def get_profile_preferred_font(profile):
    return "seuomi.ttf" if normalize_handwriting_profile(profile) == "seuomi" else "sukoya.ttf"

def get_profile_render_config(profile):
    if normalize_handwriting_profile(profile) == "seuomi":
        return {"size_mul": 0.94, "spacing_mul": 0.92, "compact_mul": 0.94}
    return {"size_mul": 1.0, "spacing_mul": 1.0, "compact_mul": 1.0}

def load_font_bytes_by_option(font_option, profile="sukoya"):
    """Return (font_bytes, actual_name). Prefer profile font without overwriting other profiles."""
    refresh_font_file_names()
    option = str(font_option or "").strip()
    if option == PROFILE_SEUOMI_VIRTUAL:
        profile = "seuomi"
        option = "seuomi.ttf"

    preferred = get_profile_preferred_font(profile)
    if preferred in font_file_names and (not option or option.startswith("__system__:") or option.lower().startswith(("sukoya", "seuomi"))):
        option = preferred

    if option.startswith("__system__:") or not option:
        system_path = resolve_system_font_path()
        if system_path:
            with open(system_path, "rb") as f:
                return f.read(), os.path.basename(system_path)
        return None, ""

    if option not in font_file_names:
        option = preferred if preferred in font_file_names else (next((f for f in font_file_names if f.lower().startswith("sukoya")), None) or (font_file_names[0] if font_file_names else ""))

    if option:
        font_path = os.path.join(font_assets_dir, option)
        if os.path.exists(font_path):
            with open(font_path, "rb") as f:
                return f.read(), option

    system_path = resolve_system_font_path()
    if system_path:
        with open(system_path, "rb") as f:
            return f.read(), os.path.basename(system_path)

    return None, ""
# sentry部分 7.7
sentry_sdk.init(
    dsn="https://ed22d5c0e3584faeb4ae0f67d19f68aa@o4505255803551744.ingest.sentry.io/4505485583253504",
    integrations=[
        StarletteIntegration(),
        FastApiIntegration(),
    ],
    # Set traces_sample_rate to 1.0 to capture 100%
    # of transactions for performance monitoring.
    # We recommend adjusting this value in production.
    traces_sample_rate=1.0,
)

# 启动计划任务线程, 定时清理
schedule_clean.start_schedule_thread()

# 创建一个logger
logger = logging.getLogger(__name__)

# 设置日志级别
logger.setLevel(logging.DEBUG)

# 创建 console handler，并设置级别为 DEBUG
ch = logging.StreamHandler()
ch.setLevel(logging.DEBUG)

# 创建 file handler，并设置级别为 DEBUG
import os
os.makedirs("logs", exist_ok=True)
fh = logging.FileHandler("logs/app.log", encoding="utf-8")
fh.setLevel(logging.DEBUG)

# 创建 formatter
formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - %(message)s")

# 把 formatter 添加到 ch 和 fh
ch.setFormatter(formatter)
fh.setFormatter(formatter)

# 把 ch 和 fh 添加到 logger
logger.addHandler(ch)
logger.addHandler(fh)

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# 自定义 422 错误响应，把字段名提取出来让前端更易读
@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    missing_fields = []
    invalid_fields = []
    for err in exc.errors():
        loc = ".".join(str(l) for l in err["loc"] if l not in ("body", "query", "path"))
        if err["type"] in ("missing", "value_error.missing"):
            missing_fields.append(loc)
        else:
            invalid_fields.append(f"{loc}: {err['msg']}")

    if missing_fields:
        message = f"缺少必填字段: {', '.join(missing_fields)}"
    else:
        message = f"字段验证失败: {', '.join(invalid_fields)}"

    return JSONResponse(
        status_code=422,
        content={
            "status": "fail",
            "message": message,
            "errors": [
                {"field": ".".join(str(l) for l in e["loc"] if l not in ("body", "query", "path")),
                 "message": e["msg"]}
                for e in exc.errors()
            ],
        },
    )

# 设置Flask app的logger级别
# app.logger.setLevel(logging.DEBUG)


SECRET_KEY = "437d75c5af744b76607fe862cf8a5a368519aca486d62c5fa69ba42c16809z88"
# app.config["SECRET_KEY"] = SECRET_KEY
# app.config["SESSION_COOKIE_SECURE"] = True
# app.config["SESSION_COOKIE_SAMESITE"] = "None"
# app.config["MAX_CONTENT_LENGTH"] = 128 * 1024 * 1024
# app.permanent_session_lifetime = timedelta(minutes=5000000)
# app.config["SESSION_TYPE"] = "filesystem"  # 设置session存储方式为文件
# Session(app)  # 初始化扩展，传入应用程序实例
limiter = Limiter(key_func=get_remote_address, default_limits=["1000 per 5 minute"])
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


# 创建一个新的白色图片，并添加间隔的线条作为背景
def create_notebook_image(
    width,
    height,
    line_spacing,
    top_margin,
    bottom_margin,
    left_margin,
    right_margin,
    font_size,
    isUnderlined,
):
    image = Image.new("RGB", (width, height), "white")

    if isUnderlined == "true":
        draw = ImageDraw.Draw(image)
        # todo  这个距离的原理不清楚7.15
        y = top_margin + line_spacing  # 开始的y坐标设为顶部边距加字体大小
        # bottom_margin -= line_spacing
        while (
            y < height - bottom_margin
        ):  # 当y坐标小于（图片高度-底部边距）时，继续画线
            draw.line((left_margin, y, width - right_margin, y), fill="black")
            y += line_spacing  # 每次循环，y坐标增加行间距
        # draw.line((left_margin, y, width - right_margin, y), fill="black")
    return image


def read_docx(file_path):
    document = Document(file_path)
    text = " ".join([paragraph.text for paragraph in document.paragraphs])
    return text


def convert_docx_to_text(docx_file_path):
    # Lightweight DOCX reader. No Pandoc download is required.
    try:
        document = Document(docx_file_path)
        parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception as e:
        logger.warning(f"DOCX read failed: {e}")
        return ""


def read_pdf(file_path):
    text = ""
    with open(file_path, "rb") as pdf_file_obj:
        pdf_reader = PyPDF2.PdfReader(pdf_file_obj)
        for page_num in range(len(pdf_reader.pages)):
            page_obj = pdf_reader.pages[page_num]
            text += page_obj.extract_text()
    return text


def handle_exceptions(f):
    @wraps(f)
    async def decorated_function(*args, **kwargs):
        try:
            return await f(*args, **kwargs)
        except Exception as e:
            logger.info("An error occurred during the request: %s", e)
            return JSONResponse({"status": "error", "message": str(e)}, status_code=500)

    return decorated_function


# WebSocket 推送（仍需内存存储，仅同进程有效）
task_websocket_connections: dict[str, set[WebSocket]] = {}
task_websocket_connections_lock = asyncio.Lock()

# 同时执行的上限——控制 generate_handwriting_impl 的真正并发数
# 超出部分在 semaphore 处排队等待，避免 CPU 密集任务互相挤占
MAX_CONCURRENT_EXECUTIONS = 2
_generation_semaphore = asyncio.Semaphore(MAX_CONCURRENT_EXECUTIONS)


def build_task_status_payload(
    task_id: str, task: Optional[GenerationTask] = None
) -> Optional[dict[str, Any]]:
    if task is None:
        task = get_generation_task(task_id)
    if task is None:
        return None
    queue_metrics = get_generation_queue_metrics(task_id)
    return {
        "status": "success",
        "task_id": task_id,
        "task_status": task.get("status"),
        "task_stage": task.get("stage"),
        "task_message": task.get("message"),
        "task_progress": task.get("progress"),
        "created_at": task.get("created_at"),
        "updated_at": task.get("updated_at"),
        "error_message": task.get("error_message"),
        "queue_pending_count": queue_metrics.get("queue_pending_count"),
        "queue_ahead_count": queue_metrics.get("queue_ahead_count"),
        "processing_count": queue_metrics.get("processing_count"),
        "active_task_count": queue_metrics.get("active_task_count"),
    }


async def register_task_websocket(task_id, websocket):
    async with task_websocket_connections_lock:
        if task_id not in task_websocket_connections:
            task_websocket_connections[task_id] = set()
        task_websocket_connections[task_id].add(websocket)


async def unregister_task_websocket(task_id, websocket):
    async with task_websocket_connections_lock:
        sockets = task_websocket_connections.get(task_id)
        if sockets is None:
            return
        sockets.discard(websocket)
        if len(sockets) == 0:
            task_websocket_connections.pop(task_id, None)


async def push_task_status_update(task_id):
    task = get_generation_task(task_id)
    if task is None:
        return
    payload = build_task_status_payload(task_id, task=task)
    if payload is None:
        return
    async with task_websocket_connections_lock:
        sockets = list(task_websocket_connections.get(task_id, set()))
    if len(sockets) == 0:
        return

    dead_sockets = []
    for socket in sockets:
        try:
            await socket.send_json(payload)
        except Exception:
            dead_sockets.append(socket)

    if dead_sockets:
        async with task_websocket_connections_lock:
            current_sockets = task_websocket_connections.get(task_id, set())
            for socket in dead_sockets:
                current_sockets.discard(socket)
            if len(current_sockets) == 0 and task_id in task_websocket_connections:
                task_websocket_connections.pop(task_id, None)


def model_to_dict(model):
    if hasattr(model, "model_dump"):
        return model.model_dump(exclude_none=True)
    return model.dict(exclude_none=True)




def stable_text_seed(text):
    value = 2166136261
    for ch in str(text or "")[:2000]:
        value ^= ord(ch)
        value = (value * 16777619) & 0xFFFFFFFF
    return value or 13579


def paste_text_glyph(page, ch, font_bytes, base_font_size, x, y, rng, page_index, fatigue_factor, chaos=0.6, ink_bleed_count=0, overlap=False):
    """Draw one character with small independent variations.
    This approximates same-character multi-shape behavior without requiring
    separate glyph image libraries.
    """
    if ch == "":
        return 0
    chaos = max(0.0, min(1.0, float(chaos)))
    try:
        ink_bleed_count = int(ink_bleed_count)
    except Exception:
        ink_bleed_count = 0
    ink_bleed_count = max(0, min(3, ink_bleed_count))
    low_punctuation = set("，。,.、；;：:！!？?")
    minus_chars = set("-−﹣－")
    is_low_punctuation = ch in low_punctuation
    is_minus = ch in minus_chars
    is_digit_symbol = bool(re.match(r"[0-9=+\-xX→∴∵≠≤≥()（）,，.。:：;/⁄₀₁₂₃₄₅₆₇₈₉⁰¹²³⁴⁵⁶⁷⁸⁹]", ch))
    size_jitter = 0.015 + chaos * 0.045
    scale_low, scale_high = (1.0 - size_jitter, 1.0 + size_jitter) if is_digit_symbol else (1.0 - size_jitter * 0.8, 1.0 + size_jitter)
    if is_minus:
        # Minus signs are visually small; keep size stable to avoid floating/oversized marks.
        scale = rng.uniform(0.98, 1.02) * fatigue_factor
    else:
        scale = rng.uniform(scale_low, scale_high) * fatigue_factor
    size = max(12, int(round(base_font_size * scale)))
    try:
        char_font = ImageFont.truetype(io.BytesIO(font_bytes), size=size)
    except Exception:
        return int(base_font_size * 0.6)

    probe = Image.new("L", (size * 3, size * 3), 0)
    d = ImageDraw.Draw(probe)
    bbox = d.textbbox((0, 0), ch, font=char_font)
    cw = max(1, bbox[2] - bbox[0])
    chh = max(1, bbox[3] - bbox[1])

    pad = max(6, int(size * 0.25))
    tile = Image.new("RGBA", (cw + pad * 2, chh + pad * 2), (0, 0, 0, 0))
    td = ImageDraw.Draw(tile)

    # Ink depth variation only. Do not auto-bold by default.
    ink_span = 4 + int(chaos * 6)
    ink = int(rng.uniform(22 - ink_span * 0.25, 23 + ink_span))
    base_x = pad - bbox[0]
    base_y = pad - bbox[1]
    td.text((base_x, base_y), ch, font=char_font, fill=(ink, ink, ink, 255), stroke_width=0)

    # Optional controlled ink bleed. 0 = clean pen, 1~3 = gradually stronger bleed.
    if ink_bleed_count > 0 and not is_low_punctuation:
        bleed_offsets = [(-1, 0), (1, 0), (0, 1), (0, -1), (-1, 1), (1, 1)]
        start_idx = rng.randint(0, len(bleed_offsets) - 1)
        for j in range(ink_bleed_count):
            ox, oy = bleed_offsets[(start_idx + j) % len(bleed_offsets)]
            alpha = 38 + j * 18
            if is_minus:
                alpha = max(22, alpha - 16)
            td.text((base_x + ox, base_y + oy), ch, font=char_font, fill=(ink, ink, ink, alpha), stroke_width=0)

    # Same-character variants are simulated by deterministic transform choices.
    angle_span = 0.10 + chaos * 0.75
    angle = rng.uniform(-angle_span, angle_span)
    if is_digit_symbol:
        angle = rng.uniform(-angle_span * 1.15, angle_span * 1.25)
    if is_low_punctuation:
        angle = rng.uniform(-0.08, 0.08)
    if is_minus:
        # Keep minus signs almost horizontal.
        angle = rng.uniform(-0.03, 0.03)
    tile = tile.rotate(angle, resample=Image.Resampling.BICUBIC, expand=True)

    base_drift_down = int(round(1 + chaos * 2))
    base_drift_up = int(round(1 + chaos * 2))
    baseline_shift = rng.randint(-base_drift_up, base_drift_down)
    x_shift = rng.uniform(-0.25 - chaos * 0.8, 0.25 + chaos * 0.8)
    if is_low_punctuation:
        # Small punctuation floated upward before. Keep it near the lower baseline.
        baseline_shift = rng.randint(0, 1)
        top = int(y + base_font_size * 0.58 + baseline_shift)
    elif is_minus:
        # Minus signs were too high because their glyph bbox is tiny.
        # Put them around mathematical midline, slightly below the old position.
        baseline_shift = rng.randint(0, 1)
        top = int(y + base_font_size * 0.37 + baseline_shift)
    else:
        top = int(y + baseline_shift + (base_font_size - size) * 0.42)
    page.alpha_composite(tile, (int(x + x_shift), top))

    extra_span_left = int(round(1 + chaos * 4))
    extra_span_right = int(round(1 + chaos * 5))
    extra_spacing = rng.randint(-extra_span_left, extra_span_right)
    if overlap and chaos > 0.35:
        extra_spacing -= rng.randint(1, max(1, int(1 + chaos * 4)))
    return max(1, int(cw + extra_spacing))


def render_sukoya_optimized_pages(text, font_bytes, data, safe_int_func):
    """Render text page-by-page with handwriting variations required by the final plan:
    - same character variation through random transforms
    - per-character spacing, size, baseline drift
    - page fatigue
    - local overlap for common terms
    - digit/symbol specific variation
    No artificial noise, blur, texture or scan shadow.
    """
    width = safe_int_func(data.get("width"), 794)
    height = safe_int_func(data.get("height"), 1123)
    line_spacing = safe_int_func(data.get("line_spacing", 52), 52)
    top_margin = safe_int_func(data.get("top_margin", 74), 74)
    bottom_margin = safe_int_func(data.get("bottom_margin", 60), 60)
    left_margin = safe_int_func(data.get("left_margin", 72), 72)
    right_margin = safe_int_func(data.get("right_margin", 60), 60)
    base_font_size = safe_int_func(data.get("font_size", 31), 31)
    is_underlined = data.get("isUnderlined", "true")
    chaos = max(0.0, min(1.0, safe_int_func(data.get("handwriting_chaos", 60), 60) / 100.0))
    ink_bleed_count = max(0, min(3, safe_int_func(data.get("ink_bleed_count", 0), 0)))
    handwriting_profile = normalize_handwriting_profile(data.get("handwriting_profile") or "sukoya")
    profile_cfg = get_profile_render_config(handwriting_profile)

    seed = stable_text_seed(text)
    rng = random.Random(seed)

    pages = []
    page_index = 0
    page = create_notebook_image(width, height, line_spacing, top_margin, bottom_margin, left_margin, right_margin, base_font_size, is_underlined).convert("RGBA")
    x = left_margin
    y = top_margin + 4

    join_terms = ("方程", "参数", "向量", "因此", "所以", "得到", "代入", "解得", "直线", "平面")
    math_symbols = set("=+-xX→∴∵≠≤≥()（）/⁄")

    def new_page():
        nonlocal page, x, y, page_index
        pages.append(page.convert("RGB"))
        page_index += 1
        page = create_notebook_image(width, height, line_spacing, top_margin, bottom_margin, left_margin, right_margin, base_font_size, is_underlined).convert("RGBA")
        x = left_margin
        y = top_margin + 4

    def new_line():
        nonlocal x, y
        x = left_margin
        y += line_spacing
        if y + line_spacing > height - bottom_margin:
            new_page()

    source = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not source.strip():
        source = " "

    i = 0
    while i < len(source):
        ch = source[i]
        if ch == "\n":
            new_line()
            i += 1
            continue

        fatigue_factor = max(0.94, 1.0 - page_index * (0.006 + chaos * 0.022))
        # Later pages become slightly more compact.
        compact_factor = max(0.88, 1.0 - page_index * (0.004 + chaos * 0.018))
        available_right = width - right_margin

        if ch == "\t":
            ch = " "

        if ch == " ":
            x += max(4, int(base_font_size * 0.32 * compact_factor)) + rng.randint(-2, 3)
            if x > available_right - base_font_size:
                new_line()
            i += 1
            continue

        # Estimate width before drawing, for wrapping.
        try:
            test_font = ImageFont.truetype(io.BytesIO(font_bytes), size=max(12, int(base_font_size * fatigue_factor)))
            bbox = ImageDraw.Draw(Image.new("L", (100, 100), 0)).textbbox((0, 0), ch, font=test_font)
            est_w = max(8, bbox[2] - bbox[0]) + 6
        except Exception:
            est_w = base_font_size

        if x + est_w > available_right:
            new_line()

        before = source[max(0, i - 1): i + 1]
        after = source[i: i + 2]
        overlap = any(term in before or term in after or source[max(0, i - 2): i + 2] in term for term in join_terms)
        if ch in math_symbols:
            overlap = False

        render_font_size = max(12, int(base_font_size * profile_cfg.get("size_mul", 1.0)))
        advance = paste_text_glyph(page, ch, font_bytes, render_font_size, x, y, rng, page_index, fatigue_factor, chaos=chaos, ink_bleed_count=ink_bleed_count, overlap=overlap)
        x += int(advance * compact_factor * profile_cfg.get("spacing_mul", 1.0) * profile_cfg.get("compact_mul", 1.0))
        i += 1

    pages.append(page.convert("RGB"))
    return pages


def post_process_sukoya_scan(image):
    """Keep output clean. Do not add noise, blur, paper texture or scan shadow."""
    try:
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")
        elif image.mode == "L":
            image = image.convert("RGB")
        # Only mild contrast/sharpness. No artificial noise/blur/texture.
        image = ImageOps.autocontrast(image, cutoff=0)
        image = ImageEnhance.Contrast(image).enhance(1.04)
        image = ImageEnhance.Sharpness(image).enhance(1.08)
        return image
    except Exception as exc:
        try:
            logger.warning("sukoya clean post-process failed: %s", exc)
        except Exception:
            pass
        return image



def estimate_content_bottom(rendered_page, blank_page, top_margin, bottom_margin, left_margin, right_margin):
    """Estimate the lowest y position occupied by handwriting by comparing
    the rendered page with a blank notebook page."""
    try:
        page = rendered_page.convert("L")
        blank = blank_page.convert("L")
        w, h = page.size
        x1 = max(0, int(left_margin))
        x2 = min(w, w - int(right_margin))
        y1 = max(0, int(top_margin))
        y2 = min(h, h - int(bottom_margin))
        last_y = y1
        min_changed = max(10, int((x2 - x1) * 0.015))
        for y in range(y1, y2):
            changed = 0
            for x in range(x1, x2):
                if abs(page.getpixel((x, y)) - blank.getpixel((x, y))) > 22:
                    changed += 1
            if changed >= min_changed:
                last_y = y
        return last_y
    except Exception:
        return int(top_margin)


def make_light_background_transparent(img):
    """Turn the diagram background transparent so notebook lines remain visible."""
    img = img.convert("RGBA")
    w, h = img.size
    samples = []
    sample_points = [(2, 2), (max(2, w - 3), 2), (2, max(2, h - 3)), (max(2, w - 3), max(2, h - 3))]
    for sx, sy in sample_points:
        if 0 <= sx < w and 0 <= sy < h:
            samples.append(img.getpixel((sx, sy))[:3])
    if samples:
        br = sum(v[0] for v in samples) / len(samples)
        bg = sum(v[1] for v in samples) / len(samples)
        bb = sum(v[2] for v in samples) / len(samples)
    else:
        br, bg, bb = (245, 245, 245)
    new_pixels = []
    for r, g, b, a in img.getdata():
        dist = abs(r - br) + abs(g - bg) + abs(b - bb)
        if dist < 34 or (r > 235 and g > 235 and b > 235):
            new_pixels.append((r, g, b, 0))
        else:
            new_pixels.append((r, g, b, a))
    img.putdata(new_pixels)
    return img


def get_small_diagram_box(page_width, page_height, avail_w=None, avail_h=None):
    """Target about 1/8 of an A4 page area for diagrams."""
    target_w = max(180, int(page_width * 0.36))
    target_h = max(120, int(page_height * 0.24))
    if avail_w is not None:
        target_w = min(target_w, int(avail_w))
    if avail_h is not None:
        target_h = min(target_h, int(avail_h))
    return target_w, target_h


def prepare_diagram_overlay(diagram_bytes, max_w, max_h):
    try:
        diag = Image.open(io.BytesIO(diagram_bytes)).convert("RGBA")
    except Exception:
        return None
    try:
        bbox = diag.getbbox()
        if bbox:
            diag = diag.crop(bbox)
    except Exception:
        pass
    diag = make_light_background_transparent(diag)
    diag.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)
    return diag


def compose_diagram_into_last_page(last_page, blank_page, diagram_bytes, line_spacing, top_margin, bottom_margin, left_margin, right_margin):
    width, height = last_page.size
    content_bottom = estimate_content_bottom(last_page, blank_page, top_margin, bottom_margin, left_margin, right_margin)
    first_line_y = int(top_margin) + int(line_spacing)
    if content_bottom < first_line_y:
        start_y = first_line_y
    else:
        relative = max(0, content_bottom - first_line_y)
        start_y = first_line_y + ((relative // max(1, int(line_spacing))) + 1) * int(line_spacing) + 6
    available_h = max(0, height - int(bottom_margin) - start_y - 6)
    available_w = max(0, width - int(left_margin) - int(right_margin) - 8)
    if available_h < 140:
        return None
    max_w, max_h = get_small_diagram_box(width, height, available_w, available_h)
    diag = prepare_diagram_overlay(diagram_bytes, max_w, max_h)
    if diag is None:
        return None
    page = last_page.convert("RGBA")
    fx = (width - diag.width) // 2
    fy = start_y
    page.alpha_composite(diag, (fx, fy))
    try:
        diag.close()
    except Exception:
        pass
    return page.convert("RGB")



def build_diagram_page(diagram_bytes, canvas_size, line_spacing, top_margin, bottom_margin, left_margin, right_margin, font=None, caption=""):
    width, height = canvas_size
    base_page = create_notebook_image(width, height, line_spacing, top_margin, bottom_margin, left_margin, right_margin, 28, "true")
    page = base_page.convert("RGBA")
    max_w, max_h = get_small_diagram_box(
        width,
        height,
        max(120, width - left_margin - right_margin - 8),
        max(120, height - top_margin - bottom_margin - 12),
    )
    diag = prepare_diagram_overlay(
        diagram_bytes,
        max_w,
        max_h,
    )
    if diag is None:
        return base_page
    fx = (width - diag.width) // 2
    fy = top_margin + line_spacing
    page.alpha_composite(diag, (fx, fy))
    try:
        diag.close()
    except Exception:
        pass
    return page.convert("RGB")

async def generate_handwriting_impl(
    base_url: str,
    params: GenerateHandwritingParams,
    background_image: Union[UploadFile, str, bytes] = File(None),
    font_file: Union[UploadFile, str, bytes] = File(None),
    progress_hook=None,
):
    def report_progress(stage, message, progress):
        if progress_hook is not None:
            progress_hook(stage=stage, message=message, progress=progress)

    report_progress("validating", "正在校验参数", 5)
    # 归一化：前端可能发字符串 "null"，需要转成 Python None
    if isinstance(background_image, str):
        background_image = None
    if isinstance(font_file, str):
        font_file = None
    # 把所有 form 字段收拢成 data dict，方便后续代码不大改
    data = model_to_dict(params)
    diagram_bytes = None
    if str(data.get("embed_diagram", "false")).lower() == "true":
        raw_diagram = data.get("diagram_png_base64") or ""
        if raw_diagram:
            try:
                diagram_bytes = base64.b64decode(raw_diagram)
            except Exception:
                diagram_bytes = None

    def safe_int(value, default=0):
        """Accept int-like or decimal strings from the frontend, e.g. '1.35'."""
        try:
            if value is None or value == "":
                return int(default)
            return int(round(float(value)))
        except (TypeError, ValueError):
            return int(default)

    report_progress("system_check", "正在检查服务器负载", 10)
    cpu_usage = psutil.cpu_percent(interval=1)  # 获取 CPU 使用率，1 秒采样间隔
    if cpu_usage > 90:
        # 如果 CPU 使用率超过 90%，返回提醒
        return JSONResponse(
            {
                "status": "waiting",
                "message": f"CPU usage is too high. Please wait and try again. current cpu_usage: {cpu_usage}%",
            },
            status_code=429,
        )  # HTTP 429: Too Many Requests
    # logger.info("已经进入generate_handwriting")
    if enable_user_auth.lower() == "true":
        # session auth 已移除，如需恢复请使用 JWT 或 cookie
        pass
    # try:
    # 先获取 form 数据
    if len(data["text"]) > 10000 and (
        base_url == "https://handwrite.sixiangjia.de/"
        or base_url == "https://handwrite.sixiangjia.de/"
    ):
        # 请自己构建应用来运行而不是使用这个网页
        return JSONResponse(
            {
                "status": "error",
                "message": "The text is too long to process. If you want to use this service, please build your own application.",
            },
            status_code=500,
        )
    # 如果存在height和width，就创建一个新的背景图     todo
    # height=int(data["height"]),
    # width=int(data["width"]),

    # 如果用户提供了宽度和高度，创建一个新的笔记本背景图像
    if "width" in data and "height" in data:
        report_progress("prepare_background", "正在创建背景图", 20)
        line_spacing = safe_int(data.get("line_spacing", 30), 30)
        top_margin = safe_int(data.get("top_margin", 0), 0)
        bottom_margin = safe_int(data.get("bottom_margin", 0), 0)
        left_margin = safe_int(data.get("left_margin", 0), 0)
        right_margin = safe_int(data.get("right_margin", 0), 0)
        width = safe_int(data["width"], 794)
        height = safe_int(data["height"], 1123)
        font_size = safe_int(data.get("font_size", 0), 31)
        isUnderlined = data.get("isUnderlined", False)
        background_image_obj = create_notebook_image(
            width,
            height,
            line_spacing,
            top_margin,
            bottom_margin,
            left_margin,
            right_margin,
            font_size,
            isUnderlined,
        )

    else:
        # 否则使用用户上传的背景图像
        report_progress("prepare_background", "正在读取背景图", 20)
        if background_image is None:
            return JSONResponse(
                {
                    "status": "fail",
                    "message": "Missing required field: background_image",
                },
                status_code=400,
            )
        if isinstance(background_image, (bytes, bytearray)):
            image_data = io.BytesIO(background_image)
        else:
            image_data = io.BytesIO(await background_image.read())

        # 使用 PIL 打开图像
        try:
            background_image_obj = Image.open(image_data)

            # 如果图像包含 Alpha 通道（模式为 'RGBA' 或 'LA'），则去除 Alpha 通道
            if background_image_obj.mode in ("RGBA", "LA"):
                # 将图像转换为 'RGB' 模式
                background_image_obj = background_image_obj.convert("RGB")

        except IOError:
            return JSONResponse(
                {"status": "error", "message": "Invalid image format"}, status_code=400
            )

    text_to_generate = data["text"]

    # Conditionally adjust spacing for English text based on user setting
    if data.get("enableEnglishSpacing", "false").lower() == "true":
        # Only apply to English words, leave Chinese text unchanged
        import re

        def replace_english_spaces(text):
            """Replace single spaces with double spaces only for English text, preserving all other whitespace"""
            # Pattern to identify English characters (including common punctuation, hyphens, underscores)
            english_pattern = r'^[a-zA-Z0-9.,!?;:\'\"()\-_]+$'

            # Split by lines to preserve newlines
            lines = text.split('\n')
            processed_lines = []

            for line in lines:
                # Only process spaces within each line, preserve tabs and other whitespace
                # Split only on spaces (not all whitespace)
                parts = line.split(' ')
                if len(parts) <= 1:
                    # No spaces in this line, keep as is
                    processed_lines.append(line)
                    continue

                result = []
                for i, part in enumerate(parts):
                    result.append(part)

                    # If this isn't the last part, check if we should add double space
                    if i < len(parts) - 1:
                        current_is_english = bool(re.match(english_pattern, part)) if part.strip() else False
                        next_is_english = bool(re.match(english_pattern, parts[i + 1])) if parts[i + 1].strip() else False

                        # Add double space only if both current and next parts are English
                        if current_is_english and next_is_english:
                            result.append('  ')  # Double space
                        else:
                            result.append(' ')   # Single space

                processed_lines.append(''.join(result))

            # Rejoin with newlines to preserve line structure
            return '\n'.join(processed_lines)

        text_to_generate = replace_english_spaces(text_to_generate)

    # if data["preview"] == "true":
    #     # 截短字符，只生成一面
    #     preview_length = 300  # 可以调整为所需的预览长度
    #     text_to_generate = text_to_generate[:preview_length]
    logger.info(f"text_to_generate: {text_to_generate}")
    # 从表单中获取字体文件并处理 7.4
    font_bytes_for_custom = None
    if font_file is not None:
        report_progress("prepare_font", "正在加载字体文件", 30)
        if isinstance(font_file, (bytes, bytearray)):
            font_bytes = font_file
        else:
            font_bytes = await font_file.read()
        font_bytes_for_custom = font_bytes
        font = ImageFont.truetype(io.BytesIO(font_bytes), size=safe_int(data["font_size"], 31))
    else:
        report_progress("prepare_font", "正在读取字体", 30)
        font_option = data.get("font_option") or SYSTEM_FONT_AUTO
        logger.info(f"font_option: {font_option}")
        logger.info(f"font_file_names: {font_file_names}")
        font_content, actual_font_name = load_font_bytes_by_option(font_option, data.get("handwriting_profile") or "sukoya")
        if font_content:
            font_bytes_for_custom = font_content
            logger.info(f"actual_font_name: {actual_font_name}")
            font = ImageFont.truetype(
                io.BytesIO(font_content), size=safe_int(data["font_size"], 31)
            )
        else:
            return JSONResponse(
                {
                    "status": "fail",
                    "message": "没有找到可用字体。把你的 sukoya.ttf 放进 ttf_files 后重启，或确认本机有中文字体。",
                },
                status_code=400,
            )

    template = Template(
        background=background_image_obj,
        font=font,
        line_spacing=safe_int(data["line_spacing"], 52),  # + font_size
        # fill=ast.literal_eval(data["fill"])[:3],  # Ignore the alpha value
        # fill=(0),#如果feel是只有一个颜色的话那么在改变墨水的时候会导致R变化而GB不变化,颜色会变红 9.17
        left_margin=safe_int(data["left_margin"], 72),
        top_margin=safe_int(data["top_margin"], 74),
        right_margin=safe_int(data["right_margin"], 60) - safe_int(data["word_spacing"], 1) * 2,
        bottom_margin=safe_int(data["bottom_margin"], 60),
        word_spacing=safe_int(data["word_spacing"], 1),
        line_spacing_sigma=float(data["line_spacing_sigma"]),  # line spacing jitter
        font_size_sigma=float(data["font_size_sigma"]),  # font size jitter
        word_spacing_sigma=float(data["word_spacing_sigma"]),  # word spacing jitter
        end_chars="，。",  # 防止特定字符因排版算法的自动换行而出现在行首
        perturb_x_sigma=float(data["perturb_x_sigma"]),  # x jitter
        perturb_y_sigma=float(data["perturb_y_sigma"]),  # y jitter
        perturb_theta_sigma=float(data["perturb_theta_sigma"]),  # 笔画旋转偏移随机扰动
        strikethrough_probability=float(
            data["strikethrough_probability"]
        ),  # 删除线概率
        strikethrough_length_sigma=float(
            data["strikethrough_length_sigma"]
        ),  # 删除线长度随机扰动
        strikethrough_width_sigma=float(
            data["strikethrough_width_sigma"]
        ),  # 删除线宽度随机扰动
        strikethrough_angle_sigma=float(
            data["strikethrough_angle_sigma"]
        ),  # 删除线角度随机扰动
        strikethrough_width=float(data["strikethrough_width"]),  # 删除线宽度
        ink_depth_sigma=float(data["ink_depth_sigma"]),  # 墨水深度随机扰动
    )

    # 创建一个BytesIO对象，用于保存.zip文件的内容
    logger.info(f"data[pdf_save]: {data['pdf_save']}")
    if not data["pdf_save"] == "true":
        report_progress("rendering", "正在生成手写图像", 45)
        # handwrite() 返回惰性 map 对象，只做文本排版（毫秒级），
        # 真正的 CPU 密集渲染在下方 for 循环消费 images 时才触发
        use_sukoya_renderer = str(data.get("use_sukoya_optimized_renderer", "true")).lower() == "true"
        if use_sukoya_renderer and font_bytes_for_custom:
            images = render_sukoya_optimized_pages(text_to_generate, font_bytes_for_custom, data, safe_int)
            logger.info("sukoya optimized renderer generated pages successfully")
        else:
            images = handwrite(text_to_generate, template)
            logger.info("handwrite initial images generated successfully")
        # 创建项目内的临时目录，避免使用系统临时目录
        project_temp_base = "./temp"
        os.makedirs(project_temp_base, exist_ok=True)
        temp_dir = tempfile.mkdtemp(dir=project_temp_base)
        unique_filename = "images_" + str(time.time())
        zip_path = f"./temp/{unique_filename}.zip"
        # 预览模式：检查是否为完整预览模式（本地开发）或单页预览模式（生产环境）
        is_preview = data["preview"] == "true"
        full_preview = data.get("full_preview", "true") if is_preview else None
        if is_preview:
            logger.info(f"Preview mode enabled, full_preview: {full_preview}")

        try:
            preview_images_base64 = []
            page_paths = []
            try:
                total_images = len(images)
                if total_images <= 0:
                    total_images = 1
            except TypeError:
                total_images = None
            page_index = 0
            for i, im in enumerate(images):
                if total_images is None:
                    dynamic_progress = min(88, 58 + min(i, 26))
                    report_progress("rendering", f"正在处理第 {i + 1} 页", dynamic_progress)
                else:
                    dynamic_progress = min(88, 58 + int((i / max(total_images,1)) * 22))
                    report_progress("rendering", f"正在处理第 {i + 1}/{total_images} 页", dynamic_progress)
                image_path = os.path.join(temp_dir, f"{page_index}.png")
                im = post_process_sukoya_scan(im)
                if safe_save_and_close_image(im, image_path):
                    logger.info(f"Image {page_index} saved successfully")
                    page_paths.append(image_path)
                else:
                    logger.error(f"Failed to save image {page_index}")
                del im
                page_index += 1

            if diagram_bytes:
                report_progress("rendering", "正在插入图示", 90)
                canvas_size = (safe_int(data.get("width"), 794), safe_int(data.get("height"), 1123))
                line_spacing = safe_int(data["line_spacing"], 52)
                top_margin = safe_int(data["top_margin"], 74)
                bottom_margin = safe_int(data["bottom_margin"], 60)
                left_margin = safe_int(data["left_margin"], 72)
                right_margin = safe_int(data["right_margin"], 60)
                blank_page = create_notebook_image(canvas_size[0], canvas_size[1], line_spacing, top_margin, bottom_margin, left_margin, right_margin, safe_int(data.get("font_size"), 31), data.get("isUnderlined", "true"))
                merged = None
                if page_paths:
                    last_path = page_paths[-1]
                    with Image.open(last_path) as lp:
                        merged = compose_diagram_into_last_page(lp.convert("RGB"), blank_page, diagram_bytes, line_spacing, top_margin, bottom_margin, left_margin, right_margin)
                    if merged is not None:
                        safe_save_and_close_image(post_process_sukoya_scan(merged), last_path)
                if merged is None:
                    diag_page = post_process_sukoya_scan(build_diagram_page(
                        diagram_bytes, canvas_size, line_spacing, top_margin, bottom_margin, left_margin, right_margin, font=font, caption=data.get("diagram_caption") or ""
                    ))
                    diag_path = os.path.join(temp_dir, f"{page_index}.png")
                    safe_save_and_close_image(diag_page, diag_path)
                    page_paths.append(diag_path)

            if is_preview:
                for image_path in page_paths:
                    with open(image_path, "rb") as f:
                        image_data = f.read()
                    if full_preview == "false":
                        safe_remove_directory(temp_dir)
                        report_progress("finalizing", "正在返回预览结果", 100)
                        return Response(content=image_data, media_type="image/png")
                    preview_images_base64.append(base64.b64encode(image_data).decode('utf-8'))
                safe_remove_directory(temp_dir)
                report_progress("finalizing", "正在返回预览结果", 100)
                return JSONResponse({"status": "success", "images": preview_images_base64})

            if not is_preview:
                report_progress("packaging", "正在打包ZIP文件", 92)
                # 创建ZIP文件
                shutil.make_archive(zip_path[:-4], "zip", temp_dir)

                # 读取ZIP文件到内存，然后立即删除文件
                try:
                    with open(zip_path, "rb") as f:
                        zip_data = f.read()

                    # 立即删除ZIP文件
                    safe_remove_file(zip_path)

                    # 从内存发送文件
                    report_progress("finalizing", "正在返回ZIP结果", 100)
                    response = Response(
                        content=zip_data,
                        media_type="application/zip",
                        headers={
                            "Content-Disposition": "attachment; filename=images.zip"
                        },
                    )
                except Exception as e:
                    logger.error(f"Failed to read ZIP file: {e}")
                    # 降级到直接读文件发送
                    with open(zip_path, "rb") as f:
                        zip_data = f.read()
                    report_progress("finalizing", "正在返回ZIP结果", 100)
                    response = Response(
                        content=zip_data,
                        media_type="application/zip",
                        headers={
                            "Content-Disposition": "attachment; filename=images.zip"
                        },
                    )
            return response
        finally:
            # 使用改进的安全删除函数
            safe_remove_directory(temp_dir)
            # ZIP文件已在上面删除，这里只是保险
    else:
        logger.info("PDF generate")
        temp_pdf_file_path = None  # 初始化变量
        report_progress("rendering", "正在生成手写图像", 45)
        use_sukoya_renderer = str(data.get("use_sukoya_optimized_renderer", "true")).lower() == "true"
        if use_sukoya_renderer and font_bytes_for_custom:
            raw_pages = render_sukoya_optimized_pages(text_to_generate, font_bytes_for_custom, data, safe_int)
        else:
            raw_pages = list(handwrite(text_to_generate, template))
        rendered_pages = [post_process_sukoya_scan(im) for im in raw_pages]
        if diagram_bytes:
            canvas_size = (safe_int(data.get("width"), 794), safe_int(data.get("height"), 1123))
            line_spacing = safe_int(data["line_spacing"], 52)
            top_margin = safe_int(data["top_margin"], 74)
            bottom_margin = safe_int(data["bottom_margin"], 60)
            left_margin = safe_int(data["left_margin"], 72)
            right_margin = safe_int(data["right_margin"], 60)
            blank_page = create_notebook_image(canvas_size[0], canvas_size[1], line_spacing, top_margin, bottom_margin, left_margin, right_margin, safe_int(data.get("font_size"), 31), data.get("isUnderlined", "true"))
            merged = None
            if rendered_pages:
                merged = compose_diagram_into_last_page(rendered_pages[-1], blank_page, diagram_bytes, line_spacing, top_margin, bottom_margin, left_margin, right_margin)
                if merged is not None:
                    rendered_pages[-1] = post_process_sukoya_scan(merged)
            if merged is None:
                rendered_pages.append(post_process_sukoya_scan(build_diagram_page(
                    diagram_bytes, canvas_size, line_spacing, top_margin, bottom_margin, left_margin, right_margin, font=font, caption=data.get("diagram_caption") or ""
                )))
        images = iter(rendered_pages)
        try:
            report_progress("packaging", "正在导出PDF文件", 92)
            temp_pdf_file_path = generate_pdf(images=images)
            # 将文件路径存储在请求上下文中，以便稍后可以访问它
            # request.temp_file_path = temp_pdf_file_path  # FastAPI Request 无此属性
            with open(temp_pdf_file_path, "rb") as f:
                pdf_data = f.read()
            report_progress("finalizing", "正在返回PDF结果", 100)
            return Response(
                content=pdf_data,
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=images.pdf"},
            )
        finally:
            # 清理生成的临时 PDF 文件
            if temp_pdf_file_path is not None and os.path.exists(temp_pdf_file_path):
                safe_remove_file(temp_pdf_file_path)
        # 2.9.2026 之前这里忘记删除temp_pdf_file_path, 根据记载是两年前就存在的问题，不知道最近为什么频繁出现
        #     if temp_pdf_file_path is not None:  # 检查变量是否已赋值
        #         for _ in range(5):  # 尝试5次
        #             try:
        #                 os.remove(temp_pdf_file_path)  # 尝试删除文件
        #                 break  # 如果成功删除，跳出循环
        #             except Exception as e:  # 捕获并处理删除文件时可能出现的异常
        #                 logger.error(f"Failed to remove temporary PDF file: {e}")
        #                 time.sleep(1)
        # unique_filename = "images_" + str(time.time()) + ".zip"

        # # 如果用户选择了保存为PDF，将所有图片合并为一个PDF文件
        # pdf_bytes = handwrite(text_to_generate, template, export_pdf=True, file_path=unique_filename)
        # logger.info("pdf generated successfully")
        # # 返回PDF文件
        # # mysql_operation(pdf_io)
        # return send_file(
        #     pdf_bytes,
        #     # attachment_filename="images.pdf",
        #     download_name="images.pdf",
        #     mimetype="application/pdf",
        #     as_attachment=True,
        # )


async def run_generation_task(task_id, base_url, payload):
    set_generation_task(
        task_id,
        status="processing",
        stage="started",
        message="任务已开始",
        progress=1,
    )
    await push_task_status_update(task_id)
    try:
        def progress_notify(**kwargs):
            set_generation_task(task_id, **kwargs)
            asyncio.create_task(push_task_status_update(task_id))

        async with _generation_semaphore:
            # 拿到信号量后，将阶段更新为"正在执行" 由于这里是协程，所以并不会占据很多资源 4.17.2026
            set_generation_task(task_id, stage="executing", message="正在生成中")
            await push_task_status_update(task_id)

            response = await generate_handwriting_impl(
                base_url=base_url,
                progress_hook=progress_notify,
                **payload,
            )
        response_body = response.body if response.body is not None else b""
        response_headers = {}
        if "content-disposition" in response.headers:
            response_headers["Content-Disposition"] = response.headers[
                "content-disposition"
            ]
        set_generation_task(
            task_id,
            status="completed",
            response_status_code=response.status_code,
            response_content_type=response.headers.get("content-type")
            or response.media_type
            or "application/octet-stream",
            response_headers=response_headers,
            response_body=response_body,
            stage="completed",
            message="任务处理完成",
            progress=100,
        )
        await push_task_status_update(task_id)
    except Exception as e:
        logger.exception("Generation task failed, task_id=%s", task_id)
        set_generation_task(
            task_id,
            status="failed",
            error_message=str(e),
            stage="failed",
            message="任务处理失败",
        )
        await push_task_status_update(task_id)


@app.post("/api/generate_handwriting")
@limiter.limit("200 per 5 minute")
@handle_exceptions  # 错误捕获的装饰器7.15
async def generate_handwriting(
    request: Request,
    background_tasks: BackgroundTasks,
    params: GenerateHandwritingParams = Depends(form_dependency_from_model(GenerateHandwritingParams)),
    background_image: Union[UploadFile, str] = File(None),
    font_file: Union[UploadFile, str] = File(None),
):
    cleanup_expired_generation_tasks()

    # ── 并发上限检查 ────────────────────────────────────────────────────
    MAX_ACTIVE_TASKS = 8  # pending + processing 总数上限，根据服务器配置调整
    # 队列满时给用户的建议等待时间（秒），固定值比瞎算可靠
    ESTIMATED_WAIT_SECONDS = 60

    active_count = get_generation_active_task_count()
    if active_count >= MAX_ACTIVE_TASKS:
        return JSONResponse(
            {
                "status": "queue_full",
                "message": "当前服务器队列已满，请稍后再试",
                "active_task_count": active_count,
                "max_active_tasks": MAX_ACTIVE_TASKS,
                "estimated_wait_seconds": ESTIMATED_WAIT_SECONDS,
            },
            status_code=503,
        )
    # ────────────────────────────────────────────────────────────────────

    background_image_bytes = None
    # 注意：starlette.datastructures.UploadFile 可能是 fastapi.UploadFile 的运行时类型
    # 用 hasattr 兼容两者：检查是否有 read 方法（UploadFile 特征）
    if hasattr(background_image, "read") and hasattr(background_image, "filename"):
        background_image_bytes = await background_image.read()

    font_file_bytes = None
    # 同样用 hasattr 兼容 starlette 和 fastapi 的 UploadFile
    if hasattr(font_file, "read") and hasattr(font_file, "filename"):
        font_file_bytes = await font_file.read()

    payload = {
        "params": params,
        "background_image": background_image_bytes,
        "font_file": font_file_bytes,
    }

    task_id = uuid4().hex
    now = time.time()
    set_generation_task(
        task_id,
        status="pending",
        stage="queued",
        message="任务排队中",
        progress=0,
        created_at=now,
        updated_at=now,
        response_status_code=None,
        response_content_type=None,
        response_headers={},
        error_message=None,
    )
    background_tasks.add_task(run_generation_task, task_id, str(request.base_url), payload)

    return JSONResponse({"status": "accepted", "task_id": task_id})


@app.websocket("/api/generate_handwriting/ws/{task_id}")
async def generate_handwriting_task_websocket(websocket: WebSocket, task_id: str):
    await websocket.accept()
    await register_task_websocket(task_id, websocket)
    try:
        task = get_generation_task(task_id)
        if task is None:
            await websocket.send_json(
                {"status": "error", "message": "Task not found", "task_id": task_id}
            )
            return
        await websocket.send_json(build_task_status_payload(task_id, task=task))

        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        pass
    finally:
        await unregister_task_websocket(task_id, websocket)


@app.get("/api/generate_handwriting/task/{task_id}")
@limiter.limit("600 per 5 minute")
@handle_exceptions
async def get_generate_handwriting_task_status(request: Request, task_id: str):
    cleanup_expired_generation_tasks()
    task = get_generation_task(task_id)
    if task is None:
        return JSONResponse(
            {"status": "error", "message": "Task not found"},
            status_code=404,
        )
    return JSONResponse(build_task_status_payload(task_id, task=task))


@app.get("/api/generate_handwriting/task/{task_id}/result")
@limiter.limit("600 per 5 minute")
@handle_exceptions
async def get_generate_handwriting_task_result(request: Request, task_id: str):
    cleanup_expired_generation_tasks()
    task = get_generation_task(task_id)
    if task is None:
        return JSONResponse(
            {"status": "error", "message": "Task not found"},
            status_code=404,
        )

    task_status = task.get("status")
    if task_status in ("pending", "processing"):
        return JSONResponse(
            {"status": "processing", "message": "Task is still running"},
            status_code=409,
        )
    if task_status == "failed":
        pop_generation_task(task_id)
        return JSONResponse(
            {"status": "error", "message": task.get("error_message", "Task failed")},
            status_code=500,
        )

    # 从磁盘文件读取响应体
    result_file_path = task.get("result_file_path")
    response_body = read_result_file(result_file_path) if result_file_path else b""
    if response_body is None:
        response_body = b""

    response = Response(
        content=response_body,
        media_type=task.get("response_content_type") or "application/octet-stream",
        status_code=task.get("response_status_code") or 200,
        headers=task.get("response_headers") or {},
    )
    pop_generation_task(task_id)
    return response


# @app.after_request
# def cleanup(response):
#     # 从请求上下文中获取文件路径
#     temp_file_path = getattr(request, 'temp_file_path', None)
#     if temp_file_path is not None:
#         # 尝试删除文件
#         try:
#             os.remove(temp_file_path)
#         except Exception as e:
#             app.logger.error(f"Failed to remove temporary PDF file: {e}")
#     # 返回原始响应
#     return response




@app.get("/api/deepseek_key_status")
@limiter.limit("60 per 5 minute")
async def deepseek_key_status(request: Request):
    """返回本机是否已保存 DeepSeek API Key。只返回状态，不返回密钥原文。"""
    key = (os.getenv("DEEPSEEK_API_KEY") or "").strip()
    masked = ""
    if key:
        masked = key[:6] + "..." + key[-4:] if len(key) > 12 else "已保存"
    return JSONResponse({"has_key": bool(key), "masked": masked, "model": os.getenv("DEEPSEEK_MODEL", "deepseek-chat")})



@app.post("/api/deepseek_clear_key")
@limiter.limit("30 per 5 minute")
async def deepseek_clear_key(request: Request):
    """Clear DeepSeek key from current process and local .env if it exists."""
    os.environ.pop("DEEPSEEK_API_KEY", None)
    os.environ.pop("SEARCH_API_KEY", None)
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    if os.path.exists(env_path):
        try:
            with open(env_path, "w", encoding="utf-8", newline="\n") as f:
                f.write("DEEPSEEK_API_KEY=\nDEEPSEEK_BASE_URL=https://api.deepseek.com\nDEEPSEEK_MODEL=deepseek-chat\n")
        except Exception as exc:
            logger.warning(f"clear key env failed: {exc}")
    return JSONResponse({"ok": True, "has_key": False, "masked": "", "model": "deepseek-chat"})



@app.post("/api/deepseek_save_key")
@limiter.limit("30 per 5 minute")
async def deepseek_save_key(request: Request):
    """把 DeepSeek API Key 保存到 backend/.env，方便下次一键启动直接用。"""
    try:
        data = await request.json()
    except Exception:
        return JSONResponse({"error": "请求内容不是有效 JSON"}, status_code=400)

    api_key = (data.get("api_key") or "").strip()
    model = (data.get("model") or "deepseek-chat").strip()
    search_key = (data.get("search_key") or "").strip()
    if model not in ["deepseek-chat", "deepseek-reasoner"]:
        model = "deepseek-chat"
    if not api_key:
        return JSONResponse({"error": "API Key cannot be empty"}, status_code=400)
    if not api_key.startswith("sk-"):
        return JSONResponse({"error": "DeepSeek API Key usually starts with sk-. Please check it."}, status_code=400)

    # Privacy mode: keep the key only in the current backend process.
    # Do not write API Key to backend/.env or any packaged file.
    os.environ["DEEPSEEK_API_KEY"] = api_key
    os.environ["DEEPSEEK_MODEL"] = model
    if search_key:
        os.environ["SEARCH_API_KEY"] = search_key

    masked = api_key[:6] + "..." + api_key[-4:] if len(api_key) > 12 else "已填写"
    return JSONResponse({"ok": True, "masked": masked, "model": model})

@app.post("/api/deepseek_polish")
@limiter.limit("30 per 5 minute")
async def deepseek_polish(request: Request):
    """DeepSeek 文本整理：用于把输入内容整理成更适合手写作业/笔记的中文。"""
    try:
        data = await request.json()
    except Exception:
        return JSONResponse({"error": "请求内容不是有效 JSON"}, status_code=400)

    text = (data.get("text") or "").strip()
    user_instruction = (data.get("instruction") or "").strip()
    api_key = (data.get("api_key") or os.getenv("DEEPSEEK_API_KEY") or "").strip()
    model = (data.get("model") or os.getenv("DEEPSEEK_MODEL") or "deepseek-chat").strip()
    if model not in ["deepseek-chat", "deepseek-reasoner"]:
        model = "deepseek-chat"

    if not text:
        return JSONResponse({"error": "请先输入要整理的文字"}, status_code=400)
    if not api_key:
        return JSONResponse({"error": "没有 DeepSeek API Key。可在启动脚本里填写，或在网页 AI 区域临时粘贴。"}, status_code=400)

    if not user_instruction:
        user_instruction = "把文字整理成适合手写作业/课堂笔记的自然中文，保留核心信息，不要使用 Markdown，不要加标题解释。"

    system_prompt = (
        "你是中文作业手写答案整理助手。输出要像学生作业本上的自然答案，别太模板化，别写免责声明。"
        "可以给出简短的解题思路，但不要暴露模型内部推理链。"
        "公式不要只输出 LaTeX 源码，能写成人类可读形式就写成人类可读形式；必要时才保留 LaTeX。"
        "不要使用 Markdown 标题、项目符号或代码块。"
    )
    payload = {
        "model": model,
        "temperature": 0.55,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"整理要求：{user_instruction}\n\n原文：\n{text}"},
        ],
    }

    req = urllib.request.Request(
        "https://api.deepseek.com/chat/completions",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Content-Type": "application/json; charset=utf-8",
            "Authorization": f"Bearer {api_key}",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
            result = json.loads(raw)
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", errors="replace")
        logger.warning(f"DeepSeek HTTP error: {detail}")
        return JSONResponse({"error": f"DeepSeek 调用失败：HTTP {e.code}"}, status_code=502)
    except Exception as e:
        logger.warning(f"DeepSeek request failed: {e}")
        return JSONResponse({"error": f"DeepSeek 调用失败：{str(e)}"}, status_code=502)

    try:
        content = result["choices"][0]["message"]["content"].strip()
    except Exception:
        return JSONResponse({"error": "DeepSeek 返回格式异常"}, status_code=502)

    return JSONResponse({"text": content})



def read_txt_file(path):
    for enc in ["utf-8", "utf-8-sig", "gb18030", "gbk"]:
        try:
            return open(path, "r", encoding=enc, errors="replace").read()
        except Exception:
            pass
    return open(path, "rb").read().decode("utf-8", errors="replace")


def try_ocr_image(path):
    """Best-effort OCR. Works when Tesseract is installed; otherwise returns a clear message."""
    try:
        import pytesseract
        img = Image.open(path)
        text = pytesseract.image_to_string(img, lang="chi_sim+eng")
        return (text or "").strip()
    except Exception as e:
        logger.warning(f"Image OCR unavailable: {e}")
        return ""


@app.post("/api/upload_extract")
@limiter.limit("120 per 5 minute")
async def upload_extract(request: Request, file: UploadFile = File(...)):
    if file is None or file.filename == "":
        return JSONResponse({"error": "No file uploaded"}, status_code=400)
    name = secure_filename(file.filename or "upload")
    suffix = os.path.splitext(name)[1].lower()
    if suffix not in [".pdf", ".docx", ".txt", ".rtf", ".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
        return JSONResponse({"error": "Only PDF, Word, TXT and image files are supported"}, status_code=400)
    tmp_dir = os.path.join(".", "upload_extract")
    os.makedirs(tmp_dir, exist_ok=True)
    path = os.path.join(tmp_dir, f"{uuid4().hex}{suffix}")
    content = await file.read()
    with open(path, "wb") as f:
        f.write(content)
    try:
        text = ""
        kind = "file"
        if suffix == ".pdf":
            kind = "pdf"
            text = read_pdf(path) or ""
        elif suffix == ".docx":
            kind = "word"
            text = convert_docx_to_text(path) or ""
        elif suffix in [".txt", ".rtf"]:
            kind = "text"
            text = read_txt_file(path) or ""
        else:
            kind = "image"
            text = try_ocr_image(path)
            if not text:
                text = "图片已上传，但当前本机没有可用 OCR。请把图片里的题目文字粘到左侧输入框，再使用 AI 生成答案。"
        return JSONResponse({"text": text.strip(), "kind": kind, "filename": file.filename})
    except Exception as e:
        logger.exception("upload extract failed")
        return JSONResponse({"error": f"File extraction failed: {str(e)}"}, status_code=500)
    finally:
        safe_remove_file(path)

@app.post("/api/textfileprocess")
@limiter.limit("200 per 5 minute")
async def textfileprocess(request: Request, file: UploadFile = File(...)):
    if file is None or file.filename == "":
        return JSONResponse({"error": "No file part in the request"}, status_code=400)

    if file and (
        file.filename.endswith(".docx")
        or file.filename.endswith(".pdf")
        or file.filename.endswith(".doc")
        or file.filename.endswith(".txt")
        or file.filename.endswith(".rtf")
    ):
        filename = secure_filename(file.filename)
        filepath = os.path.join(".", "textfileprocess", filename)  # 临时目录
        content = await file.read()
        with open(filepath, "wb") as f:
            f.write(content)
        text = "读取失败"  # Default value for text
        try:
            if file.filename.endswith(".docx"):
                text = convert_docx_to_text(filepath)
            elif file.filename.endswith(".pdf"):
                text = read_pdf(filepath)
            elif file.filename.endswith(".txt") or file.filename.endswith(".rtf"):
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif file.filename.endswith(".doc"):
                text = "doc文件暂不支持"
        except Exception as e:
            return JSONResponse(
                {"error": f"Error reading file: {str(e)}"}, status_code=500
            )

        # 删除临时文件
        safe_remove_file(filepath)

        return JSONResponse({"text": text})

    return JSONResponse({"error": "Invalid file type"}, status_code=400)


@app.post("/api/imagefileprocess")
@limiter.limit("200 per 5 minute")
async def imagefileprocess(request: Request, file: UploadFile = File(...)):
    if file is None or file.filename == "":
        return JSONResponse({"error": "No file part in the request"}, status_code=400)

    if file and (
        file.filename.endswith(".jpf")
        or file.filename.endswith(".png")
        or file.filename.endswith(".jpg")
        or file.filename.endswith(".jpeg")
    ):
        filename = secure_filename(file.filename)
        filepath = os.path.join("./imagefileprocess", filename)
        content = await file.read()
        with open(filepath, "wb") as f:
            f.write(content)
        (
            avg_l_whitespace,
            avg_r_whitespace,
            avg_t_whitespace,
            avg_b_whitespace,
            avg_distance,
        ) = identify_distance(filepath)
        safe_remove_file(filepath)
        return JSONResponse(
            {
                "marginLeft": avg_l_whitespace,
                "marginRight": avg_r_whitespace,
                "marginTop": avg_t_whitespace,
                "marginBottom": avg_b_whitespace,
                "lineSpacing": avg_distance,
            }
        )
    else:
        return JSONResponse({"error": "Invalid file type"}, status_code=400)


def get_filenames_in_dir(directory):
    refresh_font_file_names()
    files = [
        f
        for f in os.listdir(font_assets_dir)
        if os.path.isfile(os.path.join(font_assets_dir, f)) and f.lower().endswith((".ttf", ".otf"))
    ]
    return sorted(files, key=lambda x: (0 if x.lower().startswith("sukoya") else 1, x.lower()))


@app.get("/api/fonts_info")
def get_fonts_info():
    filenames = get_filenames_in_dir(font_assets_dir)
    logger.info(f"filenames: {filenames}")
    # Always provide a local system fallback so the app does not get stuck at “没有可用字体”.
    if "seuomi.ttf" not in filenames and PROFILE_SEUOMI_VIRTUAL not in filenames:
        filenames.append(PROFILE_SEUOMI_VIRTUAL)
    if SYSTEM_FONT_AUTO not in filenames:
        filenames.append(SYSTEM_FONT_AUTO)
    return JSONResponse(filenames)



@app.post("/api/upload_font")
async def upload_font(request: Request, font_file: UploadFile = File(...)):
    """Upload a personal .ttf/.otf font and refresh font list."""
    form = await request.form()
    profile = normalize_handwriting_profile(form.get("handwriting_profile") or "sukoya")
    filename = secure_filename(font_file.filename or f"{profile}.ttf")
    if not filename.lower().endswith((".ttf", ".otf")):
        return JSONResponse({"error": "只支持 .ttf / .otf 字体文件。"}, status_code=400)
    # Keep profiles isolated.
    ext = ".otf" if filename.lower().endswith(".otf") else ".ttf"
    if profile == "seuomi":
        filename = "seuomi" + ext
    elif "sukoya" not in filename.lower():
        filename = "sukoya" + ext
    os.makedirs(font_assets_dir, exist_ok=True)
    os.makedirs(font_assets_bundled_dir, exist_ok=True)
    content = await font_file.read()
    if not content:
        return JSONResponse({"error": "字体文件为空。"}, status_code=400)
    for target_dir in {font_assets_dir, font_assets_bundled_dir}:
        try:
            with open(os.path.join(target_dir, filename), "wb") as f:
                f.write(content)
        except Exception as exc:
            logger.warning(f"save uploaded font failed: {exc}")
    filenames = refresh_font_file_names()
    if SYSTEM_FONT_AUTO not in filenames:
        filenames.append(SYSTEM_FONT_AUTO)
    return JSONResponse({"ok": True, "filename": filename, "fonts": filenames})


@app.post("/api/clear_runtime_cache")
def clear_runtime_cache():
    """Clear backend temporary files. Does not delete DeepSeek key or fonts."""
    cleared = []
    for rel in ["temp", "textfileprocess", "imagefileprocess"]:
        path = os.path.abspath(rel)
        try:
            if os.path.isdir(path):
                shutil.rmtree(path, ignore_errors=True)
            os.makedirs(path, exist_ok=True)
            cleared.append(rel)
        except Exception as exc:
            logger.warning(f"clear cache failed for {rel}: {exc}")
    return JSONResponse({"ok": True, "cleared": cleared})


def mysql_operation(image_data):
    # cursor = current_app.cnx.cursor()
    # username = session["username"]
    username = None  # session 已移除
    # 先检查用户是否已存在
    # cursor.execute("SELECT * FROM user_images WHERE username=%s", (username,))
    # result = cursor.fetchone()

    # 根据查询结果来判断应该插入新纪录还是更新旧纪录
    # if result is None:
    #     # 如果用户不存在，插入新纪录
    #     sql = "INSERT INTO user_images (username, image) VALUES (%s, %s)"
    #     params = (username, image_data)
    # else:
    #     # 如果用户已存在，更新旧纪录
    #     sql = "UPDATE user_images SET image=%s WHERE username=%s"
    #     params = (image_data, username)
    try:
        pass
        # 执行 SQL 语句
        # 提交到数据库执行
        # cursor.execute(sql, params)
        # current_app.cnx.commit()
    except Exception as e:
        # 发生错误时回滚
        # current_app.cnx.rollback()
        logger.info(f"An error occurred: {e}")


# @app.route("/api/login", methods=["POST"])
# def login():
#     data = request.get_json()
#     username = data.get("username")
#     password = data.get("password")
#     logger.info(f"Received username: {username}")  # 打印接收到的用户名
#     logger.info(f"Received password: {password}")  # 打印接收到的密码
#     try:
#         cursor = current_app.cnx.cursor()
#         cursor.execute(
#             f"SELECT password FROM user_images WHERE username=%s", (username,)
#         )
#         result = cursor.fetchone()
#     except Exception as e:
#         logger.error(f"An error occurred: {e}")
#         return jsonify({"error": "An error occurred"}), 500

#     if result and result[0] == password:
#         session["username"] = username
#         session.permanent = True
#         logger.info(f"Login success for user: {username}")
#         return {"status": "success"}, 200
#     else:
#         logger.error(f"Login failed for user: {username}")
#         return {
#             "status": "failed",
#             "error": "Login failed. Check your username and password.",
#         }, 401


# @app.route("/api/register", methods=["POST"])
# def register():
#     data = request.get_json()
#     username = data.get("username")
#     password = data.get("password")
#     try:
#         cursor = current_app.cnx.cursor()
#         cursor.execute(f"SELECT * FROM user_images WHERE username=%s", (username,))
#         result = cursor.fetchone()
#     except Exception as e:
#         logger.error(f"An error occurred: {e}")
#         return jsonify({"error": "An error occurred"}), 500

#     if not result:
#         try:
#             cursor.execute(
#                 f"INSERT INTO user_images (username, password) VALUES (%s, %s)",
#                 (username, password),
#             )
#             current_app.cnx.commit()
#             session["username"] = username
#             logger.info(f"User: {username} registered successfully.")
#             return jsonify(
#                 {
#                     "status": "success",
#                     "message": "Account created successfully. You can now log in.",
#                 }
#             )
#         except mysql.connector.Error as err:
#             logger.error(f"Failed to insert user: {username} into DB. Error: {err}")
#             return (
#                 jsonify(
#                     {
#                         "status": "fail",
#                         "message": "Error occurred during registration.",
#                     }
#                 ),
#                 500,
#             )
#     else:
#         logger.error(f"Username: {username} already exists.")
#         return (
#             jsonify(
#                 {
#                     "status": "fail",
#                     "message": "Username already exists. Choose a different one.",
#                 }
#             ),
#             400,
#         )


# 捕获所有未捕获的异常，返回给前端，只能用于生产环境7.12
# @app.errorhandler(Exception)
# def handle_exception(e):
#     # Pass the error to Flask's default error handling.
#       tb = traceback.format_exception(etype=type(e), value=e, tb=e.__traceback__)
#     response = {
#
#             "type": type(e).__name__,  # The type of the exception
#             "message": str(e),  # The message of the exception
#
#     }
#     return jsonify(response), 500


# @app.before_request
# def before_request():
#     if enable_user_auth.lower() == "true":
#         current_app.cnx = mysql.connector.connect(
#             host=mysql_host, user="myuser", password="mypassword", database="mydatabase"
#         )
#     else:
#         pass


@app.middleware("http")
async def after_request(request: Request, call_next):
    response = await call_next(request)
    if enable_user_auth.lower() == "true":
        # if hasattr(current_app, "cnx"):
        #     current_app.cnx.close()
        # 仅用于调试 7.13
        # session.clear()
        return response
    else:
        print(response)
        return response


if __name__ == "__main__":
    import uvicorn

    # 启动时清理之前标记的目录
    cleanup_marked_directories()
    uvicorn.run(app, host="0.0.0.0", port=5005)


# poetry
def main():
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=5005)

    # good luck 6/16/2023
    # thank you 2/14/2025


"""    
数据库初始化操作

CREATE TABLE user_images (
    id INT AUTO_INCREMENT PRIMARY KEY,
    username VARCHAR(255) UNIQUE, 
    password VARCHAR(255), 
    image BLOB,
    upload_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);


数据库结构
mysql -u root -p进入数据库
USE your_database;数据库中的一个库
describe user_images;表：
"""
